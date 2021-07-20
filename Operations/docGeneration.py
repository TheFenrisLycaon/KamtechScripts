import docx
from docx.shared import Inches
import os
import shutil
from pdf2image import convert_from_path

POPPLER_PATH = './utils/bin'


def createStatement():
    # path to the pdfs of financial statement
    financialStatements = './financialStatements/'
    temp = './.cacheFin/'  # temporary directory that is created and deleted automatically

    pages = []
    for i in os.listdir(financialStatements):
        try:
            pages.extend(convert_from_path(financialStatements+i))
        except:
            pages.extend(convert_from_path(
                financialStatements+i, poppler_path=POPPLER_PATH))

        os.makedirs(temp, exist_ok=True)

        for j in range(len(pages)):
            pages[j].save('{}/{:03d}.jpg'.format(temp, j+1), 'JPEG')

        # creating doc from generated Images
        doc = docx.Document()

        for k in os.listdir(temp):
            doc.add_picture(temp+k, width=Inches(5))
            doc.add_page_break()

        # path to save the generated document
        doc.save(f'./{i.split(".")[0]}.docx')
        shutil.rmtree(temp)


def createCertifiates():
    certi = './Certificates/'  # path to certificates directory
    temp = './.cacheCer/'  # temporary directory that is created and deleted automatically

    pages = []
    for i in os.listdir(certi):
        try:
            pages.extend(convert_from_path(certi+i))
        except:
            print(1)
            pages.extend(convert_from_path(certi+i,
                                           poppler_path=POPPLER_PATH))

    os.makedirs(temp, exist_ok=True)

    for i in range(len(pages)):
        pages[i].save('{}/{:03d}.jpg'.format(temp, i+1), 'JPEG')

    # creating doc from generated Images
    doc = docx.Document()

    for i in os.listdir(temp):
        doc.add_picture(temp+i, width=Inches(5))
        doc.add_page_break()

    doc.save('./Certifiates.docx')  # path to save the generated document
    shutil.rmtree(temp)


def createLicense():
    # path to business license image directory
    businessImgs = './BusinessLicences/'
    doc = docx.Document()

    for img in os.listdir(businessImgs):
        doc.add_picture(businessImgs+img, width=Inches(5))
        doc.add_page_break()

    doc.save('./Licences.docx')  # path to output file


def createWork():
    pathWork = sorted(['./workExp/work/'+ x for x in os.listdir('./workExp/work/')])
    pathComp = sorted(['./workExp/comp/' + x for x in os.listdir('./workExp/comp/')])
    temp = './.cacheWork/'  # temporary directory that is created and deleted automatically

    for i in range(len(pathWork)):
        try:
            pagesWork = convert_from_path(pathWork[i])
            pagesComp = convert_from_path(pathComp[i])
        except:
            pagesWork = convert_from_path(pathWork[i], poppler_path=POPPLER_PATH)
            pagesComp = convert_from_path(pathWork[i], poppler_path=POPPLER_PATH)

        os.makedirs(temp, exist_ok=True)

        for j in range(len(pagesWork)):
            pagesWork[j].save(
                '{}/{}_0_{:03d}.jpg'.format(temp, i, j+1), 'JPEG')
        
        for j in range(len(pagesComp)):
            pagesComp[j].save(
                '{}/{}_1_{:03d}.jpg'.format(temp, i, j+1), 'JPEG')


    doc = docx.Document()

    for img in sorted(os.listdir(temp)):
        doc.add_picture(temp+img, width=Inches(5))
        doc.add_page_break()

    # doc.save(f'./{tender_folder}/work_pdfs.docx')
    doc.save(f'./work_pdfs.docx')

    shutil.rmtree(temp)

def conv(pdfPath, outPath):
    print('Satrted')
    # pages = convert_from_path(pdfPath)
    os.makedirs(outPath, exist_ok=True)

    count = 0
    print('saving')
    for i in convert_from_path(pdfPath):
        i.save('{}/{:03d}.jpg'.format(outPath, count+1), 'JPEG')
        count+=1
    print('done')


conv('./Sample4.pdf',
     './Compression/Org')

conv('./temp4_0.pdf',
     './Compression/Test0')

conv('./temp4_1.pdf',
     './Compression/Test1')

conv('./temp4_2.pdf',
     './Compression/Test2')

conv('./temp4_3.pdf',
     './Compression/Test3')

conv('./temp4_4.pdf',
     './Compression/Test4')
# createStatement()
# createLicense()
# createCertifiates()
# createWork()
