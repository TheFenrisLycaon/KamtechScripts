# Import docx NOT python-docx
import docx
from docx.shared import Inches

def insertImg(doc, title=None, imgPath, sizeX, sizeY):
    """
    doc :: The document object
    ttile :: Heading to be added
    imgPath :: Path to the image to be added
    sizeX, sizeY :: Width and heigth respectively
    """
    doc.add_heading(title, 0)
    doc.add_picture(imgPath, width=Inches(sizeX), height=Inches(sizeY))
    doc.add_page_break()


# Create an instance of a word document
doc = docx.Document()

insertImg(doc, 'Business License', r'./inserting images/logo.jpg', 5, 7)
insertImg(doc, 'GST', r'./inserting images/gst.jpg', 5, 7)

doc.save('./inserting images/out.docx')

# ! LEGACY CODE 
# # Add a Title to the document
# doc.add_heading('Business License', 0)
# doc.add_picture('inserting images/logo.jpg',
#                 width=Inches(2), height=Inches(2))

# doc.add_page_break()

# # inserting 2 image
# doc.add_heading('GST', 0)
# doc.add_picture('inserting images/gst.jpg', 
#                 width=Inches(2), height=Inches(2))
# # Now save the document to a location
# doc.save('inserting images/images.docx')
